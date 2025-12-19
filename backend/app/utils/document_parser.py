"""Document parser utilities for extracting text from uploaded files."""

import io
import logging
import re
import unicodedata
from pathlib import Path

logger = logging.getLogger(__name__)

# Vietnamese character mapping to ASCII
VIETNAMESE_MAP = {
    'à': 'a', 'á': 'a', 'ả': 'a', 'ã': 'a', 'ạ': 'a',
    'ă': 'a', 'ằ': 'a', 'ắ': 'a', 'ẳ': 'a', 'ẵ': 'a', 'ặ': 'a',
    'â': 'a', 'ầ': 'a', 'ấ': 'a', 'ẩ': 'a', 'ẫ': 'a', 'ậ': 'a',
    'đ': 'd',
    'è': 'e', 'é': 'e', 'ẻ': 'e', 'ẽ': 'e', 'ẹ': 'e',
    'ê': 'e', 'ề': 'e', 'ế': 'e', 'ể': 'e', 'ễ': 'e', 'ệ': 'e',
    'ì': 'i', 'í': 'i', 'ỉ': 'i', 'ĩ': 'i', 'ị': 'i',
    'ò': 'o', 'ó': 'o', 'ỏ': 'o', 'õ': 'o', 'ọ': 'o',
    'ô': 'o', 'ồ': 'o', 'ố': 'o', 'ổ': 'o', 'ỗ': 'o', 'ộ': 'o',
    'ơ': 'o', 'ờ': 'o', 'ớ': 'o', 'ở': 'o', 'ỡ': 'o', 'ợ': 'o',
    'ù': 'u', 'ú': 'u', 'ủ': 'u', 'ũ': 'u', 'ụ': 'u',
    'ư': 'u', 'ừ': 'u', 'ứ': 'u', 'ử': 'u', 'ữ': 'u', 'ự': 'u',
    'ỳ': 'y', 'ý': 'y', 'ỷ': 'y', 'ỹ': 'y', 'ỵ': 'y',
    'À': 'A', 'Á': 'A', 'Ả': 'A', 'Ã': 'A', 'Ạ': 'A',
    'Ă': 'A', 'Ằ': 'A', 'Ắ': 'A', 'Ẳ': 'A', 'Ẵ': 'A', 'Ặ': 'A',
    'Â': 'A', 'Ầ': 'A', 'Ấ': 'A', 'Ẩ': 'A', 'Ẫ': 'A', 'Ậ': 'A',
    'Đ': 'D',
    'È': 'E', 'É': 'E', 'Ẻ': 'E', 'Ẽ': 'E', 'Ẹ': 'E',
    'Ê': 'E', 'Ề': 'E', 'Ế': 'E', 'Ể': 'E', 'Ễ': 'E', 'Ệ': 'E',
    'Ì': 'I', 'Í': 'I', 'Ỉ': 'I', 'Ĩ': 'I', 'Ị': 'I',
    'Ò': 'O', 'Ó': 'O', 'Ỏ': 'O', 'Õ': 'O', 'Ọ': 'O',
    'Ô': 'O', 'Ồ': 'O', 'Ố': 'O', 'Ổ': 'O', 'Ỗ': 'O', 'Ộ': 'O',
    'Ơ': 'O', 'Ờ': 'O', 'Ớ': 'O', 'Ở': 'O', 'Ỡ': 'O', 'Ợ': 'O',
    'Ù': 'U', 'Ú': 'U', 'Ủ': 'U', 'Ũ': 'U', 'Ụ': 'U',
    'Ư': 'U', 'Ừ': 'U', 'Ứ': 'U', 'Ử': 'U', 'Ữ': 'U', 'Ự': 'U',
    'Ỳ': 'Y', 'Ý': 'Y', 'Ỷ': 'Y', 'Ỹ': 'Y', 'Ỵ': 'Y',
}


def sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe storage."""
    if not filename:
        return "unnamed_file"
    
    # Separate name and extension
    path = Path(filename)
    name = path.stem
    ext = path.suffix.lower()
    
    # Convert Vietnamese to ASCII
    result = []
    for char in name:
        if char in VIETNAMESE_MAP:
            result.append(VIETNAMESE_MAP[char])
        else:
            result.append(char)
    name = ''.join(result)
    
    # Replace spaces and special chars with underscore
    name = re.sub(r'[\s\-]+', '_', name)  # spaces and hyphens to underscore
    name = re.sub(r'[^\w]', '', name)  # remove non-alphanumeric (except underscore)
    name = re.sub(r'_+', '_', name)  # collapse multiple underscores
    name = name.strip('_')  # remove leading/trailing underscores
    
    # Ensure name is not empty
    if not name:
        name = "document"
    
    # Limit filename length to avoid filesystem issues
    if len(name) > 100:
        name = name[:100]
    
    safe_name = f"{name}{ext}"
    logger.info(f"[DocumentParser] Sanitized filename: '{filename}' -> '{safe_name}'")
    
    return safe_name


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from Word document (.docx).
    
    Args:
        file_bytes: Raw bytes of the .docx file
        
    Returns:
        Extracted text content
    """
    from docx import Document
    
    doc = Document(io.BytesIO(file_bytes))
    
    paragraphs = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))
    
    extracted_text = "\n\n".join(paragraphs)
    logger.info(f"[DocumentParser] Extracted {len(extracted_text)} characters from .docx")
    
    return extracted_text


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text file (.txt).
    
    Args:
        file_bytes: Raw bytes of the .txt file
        
    Returns:
        Text content
    """
    # Try UTF-8 first, then fall back to other encodings
    encodings = ['utf-8', 'utf-8-sig', 'cp1252', 'latin-1']
    
    for encoding in encodings:
        try:
            text = file_bytes.decode(encoding)
            logger.info(f"[DocumentParser] Decoded .txt with {encoding} ({len(text)} chars)")
            return text
        except UnicodeDecodeError:
            continue
    
    # Last resort: decode with errors ignored
    text = file_bytes.decode('utf-8', errors='ignore')
    logger.warning(f"[DocumentParser] Decoded .txt with utf-8 (errors ignored)")
    return text


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Auto-detect file type and extract text.
    
    Args:
        filename: Original filename with extension
        file_bytes: Raw bytes of the file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file type is not supported
    """
    ext = Path(filename).suffix.lower()
    
    if ext == ".docx":
        return extract_text_from_docx(file_bytes)
    elif ext == ".txt":
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .docx, .txt")
